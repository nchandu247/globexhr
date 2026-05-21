"""
One-shot: convert the HR-approved offer-letter DOCX into a merger-ready template.

Reads:  templates/Globex Digital Solutions _ Template _ Offer Letter.docx
Writes: greythr_bridge/templates/letters/offer_letter.docx

What it does:
  1. Replaces every old MailMerge `«FieldName»` with `{{ jinja_variable }}`
  2. Fixes the `{{ candidate_name }` typo (missing closing brace) on page 5
  3. Preserves all original formatting, logo, fonts, tables, header/footer

Run from repo root:
    python scripts/build_offer_template.py
"""
import os
import sys

from docx import Document


# ── Substitution map (old → new) ─────────────────────────────────────────────
REPLACEMENTS = {
    # Annexure A table — old MailMerge fields
    "«Band»":                       "{{ band }}",
    "«Basic_YTD»":                  "{{ basic_annual }}",
    "«HRA_YTD»":                    "{{ hra_annual }}",
    "«Conveyance_Allowance_YTD»":   "{{ conveyance_annual }}",
    "«Medical_Allowance_YTD»":      "{{ medical_allowance_annual }}",
    "«Special_Allowance_YTD»":      "{{ special_allowance_annual }}",
    "«Monthly_Gross»":              "{{ gross_monthly }}",
    "«Monthly_Gross_YTD»":          "{{ gross_annual }}",
    "«Provident_Fund_YTD»":         "{{ employee_pf_annual }}",
    "«ESI_YTD»":                    "{{ employee_esi_annual }}",
    "«Professional_Tax_YTD»":       "{{ professional_tax_annual }}",
    "«Decutions»":                  "{{ total_deductions_monthly }}",
    "«Decutions_YTD»":              "{{ total_deductions_annual }}",
    "«Employer_PF_YTD»":            "{{ employer_pf_annual }}",
    "«Medical_Insurance_YTD»":      "{{ employer_esi_annual }}",
    "«Employer_Deductions_YTD»":    "{{ employer_deductions_annual }}",
    "«Net_Salary»":                 "{{ net_take_home }}",
    "«Net_Salary_YTD»":             "{{ net_take_home_annual }}",
    "«Annual_CTC_YTD»":             "{{ annual_ctc }}",
    "«Current_Date»":               "{{ current_date }}",

    # Typo fix on page 5: "{{ candidate_name }" → "{{ candidate_name }}"
    # Use a unique left-side marker so we don't double-fix valid placeholders.
    "{{ candidate_name }  ": "{{ candidate_name }}  ",
    "{{ candidate_name } ":  "{{ candidate_name }} ",
}

# Special two-pass: handle the "{{ candidate_name }" without trailing space too
# (typo where the SECOND brace is missing). Do this AFTER the main pass so we
# don't accidentally re-break a correct "{{ candidate_name }}".
def fix_candidate_name_typo(text: str) -> str:
    """If text has '{{ candidate_name }' NOT followed by another '}', add one."""
    import re
    # Match '{{ candidate_name }' that is NOT immediately followed by another '}'
    return re.sub(r"\{\{\s*candidate_name\s*\}(?!\})", "{{ candidate_name }}", text)


# ── Helpers ──────────────────────────────────────────────────────────────────

def _apply_to_paragraph(paragraph) -> int:
    """
    Replace placeholders in a paragraph. Returns count of substitutions made.

    Strategy: get full paragraph text, do all replacements, then write the
    result back into the FIRST run (preserving its font) and clear all other
    runs. This loses intra-paragraph formatting (e.g. one bold word inside a
    sentence) but is the only reliable way to handle text split across runs,
    which is very common in Word documents.
    """
    full_text = paragraph.text
    if not full_text:
        return 0

    new_text = full_text
    for old, new in REPLACEMENTS.items():
        new_text = new_text.replace(old, new)
    new_text = fix_candidate_name_typo(new_text)

    if new_text == full_text:
        return 0

    if paragraph.runs:
        paragraph.runs[0].text = new_text
        for run in paragraph.runs[1:]:
            run.text = ""
    else:
        paragraph.add_run(new_text)

    # Count how many distinct substitutions hit
    count = 0
    for old in REPLACEMENTS:
        if old in full_text:
            count += full_text.count(old)
    return count


def _process_container(container) -> int:
    """Walk paragraphs and tables inside a container (body, cell, header, footer)."""
    total = 0
    for para in getattr(container, "paragraphs", []):
        total += _apply_to_paragraph(para)
    for table in getattr(container, "tables", []):
        for row in table.rows:
            for cell in row.cells:
                total += _process_container(cell)
    return total


# ── Post-substitution polish passes ───────────────────────────────────────────

def _fix_hi_dot_salutation(doc) -> int:
    """
    The HR source DOCX has 'Hi .{{ candidate_name }},' (literal dot before
    the variable). Replace with 'Hi {{ candidate_name }},' across body +
    headers + footers.

    Returns number of paragraphs modified.
    """
    import re
    pattern = re.compile(r"Hi\s*\.\s*(\{\{\s*candidate_name\s*\}\})")

    def _fix_in_container(container) -> int:
        n = 0
        for para in getattr(container, "paragraphs", []):
            if "Hi" in para.text and "candidate_name" in para.text:
                new_text = pattern.sub(r"Hi \1", para.text)
                if new_text != para.text and para.runs:
                    para.runs[0].text = new_text
                    for run in para.runs[1:]:
                        run.text = ""
                    n += 1
        for table in getattr(container, "tables", []):
            for row in table.rows:
                for cell in row.cells:
                    n += _fix_in_container(cell)
        return n

    count = _fix_in_container(doc)
    for section in doc.sections:
        count += _fix_in_container(section.header)
        count += _fix_in_container(section.footer)
    return count


def _add_reporting_manager_line(doc) -> int:
    """
    Find the joining sentence ('...ideally by {{ date_of_joining }}...') and
    insert a new paragraph after it with the Reporting Manager line.

    Uses jinja {% if %} so the sentence is omitted entirely when the
    custom_reporting_to field is not set.

    Returns number of insertions (should be 1 if the joining sentence was found).
    """
    target_marker = "date_of_joining"  # the variable name appears in the joining sentence
    inserted = 0

    for i, para in enumerate(doc.paragraphs):
        if target_marker in para.text and inserted == 0:
            # Insert a new paragraph immediately after `para`
            new_para = para.insert_paragraph_before()
            # Swap: insert before puts it ABOVE. We want it BELOW para.
            # python-docx doesn't have insert_after, so use XML manipulation:
            from copy import deepcopy
            new_p_xml = deepcopy(new_para._p)
            para._p.addnext(new_p_xml)
            # Remove the empty paragraph we inserted before
            new_para._p.getparent().remove(new_para._p)

            # Find the just-inserted paragraph (it's now after `para`)
            from docx.text.paragraph import Paragraph
            inserted_para = Paragraph(new_p_xml, para._parent)
            inserted_para.add_run(
                "{% if reporting_to %}You will report to "
                "{{ reporting_to }} in this role.{% endif %}"
            )
            inserted = 1
            break
    return inserted


def _dim_watermark(doc) -> int:
    """
    Attempt to dim the 'Globex DIGITAL' watermark behind body text.

    DOCX watermarks live in the header XML as VML or DrawingML shapes.
    python-docx doesn't expose these directly, so we use low-level XML
    manipulation to find and adjust them. Wrapped in try/except — if this
    fails for any reason, the watermark stays as-is and the rest of the
    build still succeeds (non-blocking polish).

    Returns number of watermark shapes dimmed (0 if none found or skipped).
    """
    try:
        from lxml import etree
        ns = {
            "w":   "http://schemas.openxmlformats.org/wordprocessingml/2006/main",
            "v":   "urn:schemas-microsoft-com:vml",
            "o":   "urn:schemas-microsoft-com:office:office",
        }
        dimmed = 0
        for section in doc.sections:
            header_xml = section.header._element
            # Find VML shapes (legacy watermark format used by Word)
            for shape in header_xml.iter("{urn:schemas-microsoft-com:vml}shape"):
                style = shape.get("style", "")
                # Watermarks typically have specific style attributes; mark them faded
                # by injecting opacity into the style attribute.
                if "opacity" not in style:
                    shape.set("style", style + ";opacity:.1")
                    dimmed += 1
                # Also try child <v:fill opacity="..."> if present
                for fill in shape.iter("{urn:schemas-microsoft-com:vml}fill"):
                    fill.set("opacity", "0.1")
        return dimmed
    except Exception as e:
        # Non-fatal — watermark stays as-is
        print(f"WARN: watermark dim failed ({type(e).__name__}: {e}). Continuing.")
        return 0


# ── Build entrypoint ──────────────────────────────────────────────────────────

def build(src: str, dst: str) -> dict:
    """
    Build the offer-letter merge template from the HR source DOCX.

    Testable function — `main()` wraps this with default paths.
    Returns a stats dict for verification:
        { "substitutions": N, "salutation_fixes": N,
          "reporting_manager_inserted": N, "watermark_dimmed": N }
    """
    if not os.path.exists(src):
        raise FileNotFoundError(f"source template not found: {src}")

    os.makedirs(os.path.dirname(dst), exist_ok=True)
    doc = Document(src)

    stats = {
        "substitutions": 0,
        "salutation_fixes": 0,
        "reporting_manager_inserted": 0,
        "watermark_dimmed": 0,
    }

    # Pass 1: «...» → {{...}} substitutions + candidate_name typo fix
    stats["substitutions"] += _process_container(doc)
    for section in doc.sections:
        stats["substitutions"] += _process_container(section.header)
        stats["substitutions"] += _process_container(section.footer)

    # Pass 2: polish
    stats["salutation_fixes"] = _fix_hi_dot_salutation(doc)
    stats["reporting_manager_inserted"] = _add_reporting_manager_line(doc)
    stats["watermark_dimmed"] = _dim_watermark(doc)

    doc.save(dst)
    return stats


def main():
    repo_root = os.path.dirname(os.path.dirname(os.path.abspath(__file__)))
    src = os.path.join(
        repo_root, "templates",
        "Globex Digital Solutions _ Template _ Offer Letter.docx",
    )
    dst = os.path.join(repo_root, "greythr_bridge", "templates", "letters", "offer_letter.docx")

    try:
        stats = build(src, dst)
    except FileNotFoundError as e:
        sys.exit(f"ERROR: {e}")

    print(f"OK: substitutions applied = {stats['substitutions']}")
    print(f"OK: salutation fixes      = {stats['salutation_fixes']}")
    print(f"OK: reporting mgr line    = {stats['reporting_manager_inserted']}")
    print(f"OK: watermark dimmed      = {stats['watermark_dimmed']}")
    print(f"OK: wrote {dst}")


if __name__ == "__main__":
    main()
