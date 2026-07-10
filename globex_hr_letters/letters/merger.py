"""
Render primitives for Globex HR letters.

Two render engines, selected per Letter Type:

  HTML  — Jinja2 template under templates/letters/html/ (extends _base.html
          for letterhead, watermark, brand CSS) rendered to PDF via
          WeasyPrint. Pixel-perfect; used by the shipped template library.
          Zoho Sign text tags are embedded by the base template when the
          context sets requires_signature.

  DOCX  — HR-uploaded Word file with {{placeholder}} variables rendered via
          docxtpl. For signature letters the DOCX (with appended Zoho text
          tags) is uploaded to Zoho Sign directly — Zoho converts to PDF
          server-side, so no LibreOffice needed. For plain letters the DOCX
          is converted to PDF via LibreOffice when available.

This module does rendering only — no Frappe document I/O beyond template
file resolution. Context building lives in letters/engine.py.
"""
import os
import tempfile

import frappe
from docxtpl import DocxTemplate

from .pdf_convert import docx_to_pdf_bytes

_HTML_TEMPLATE_DIR = os.path.normpath(
    os.path.join(os.path.dirname(__file__), "..", "templates", "letters", "html")
)


# ── template discovery ────────────────────────────────────────────────────────

def html_template_exists(template_filename: str) -> bool:
    """True if the named template exists under templates/letters/html/."""
    if not template_filename or "/" in template_filename or "\\" in template_filename:
        return False
    return os.path.exists(os.path.join(_HTML_TEMPLATE_DIR, template_filename))


def scan_html_placeholders(template_filename: str) -> set:
    """
    Return the set of undeclared Jinja2 variables in an HTML template
    (including variables used by the base template it extends).
    """
    from jinja2 import Environment, FileSystemLoader, meta

    env = Environment(loader=FileSystemLoader(_HTML_TEMPLATE_DIR))
    names = set()
    # Walk the template plus anything it extends/includes so base-template
    # variables (company_name, signatory_name, ...) are captured too.
    pending = [template_filename]
    seen = set()
    while pending:
        current = pending.pop()
        if current in seen:
            continue
        seen.add(current)
        source = env.loader.get_source(env, current)[0]
        ast = env.parse(source)
        names |= meta.find_undeclared_variables(ast)
        pending.extend(ref for ref in meta.find_referenced_templates(ast) if ref)
    return names


def scan_docx_placeholders(docx_path: str) -> set:
    """Return the set of undeclared Jinja2 variables in a DOCX template."""
    tpl = DocxTemplate(docx_path)
    return tpl.get_undeclared_template_variables()


# ── HTML engine ───────────────────────────────────────────────────────────────

def merge_to_pdf_via_html(template_filename: str, context: dict) -> bytes:
    """
    Render an HTML template (Jinja2) and produce PDF bytes via WeasyPrint.

    Templates live in templates/letters/html/. Each letter template extends
    _base.html (letterhead, watermark, footer, Zoho tag area) and uses
    _styles.css for typography, layout, and brand colours.

    When context["requires_signature"] is truthy the base template renders
    invisible Zoho Sign text tags ({{S:R1*}} and {{S:R2*}}) so Zoho
    auto-creates Signature fields on upload — no field coordinates needed.

    Raises frappe.ValidationError if the template file is missing.
    Raises RuntimeError if WeasyPrint rendering fails.
    """
    from weasyprint import HTML, CSS

    template_path = os.path.join(_HTML_TEMPLATE_DIR, template_filename)
    css_path = os.path.join(_HTML_TEMPLATE_DIR, "_styles.css")

    if not os.path.exists(template_path):
        frappe.throw(
            f"HTML template not found: {template_filename}<br>"
            f"Expected path: {template_path}",
            title="Template Missing",
        )

    # Use Jinja2 directly so we can resolve {% extends '_base.html' %} against
    # the html/ directory. frappe.render_template doesn't know about our path.
    from jinja2 import Environment, FileSystemLoader, select_autoescape
    env = Environment(
        loader=FileSystemLoader(_HTML_TEMPLATE_DIR),
        autoescape=select_autoescape(["html", "xml"]),
    )
    template = env.get_template(template_filename)
    rendered_html = template.render(**context)

    pdf_bytes = HTML(string=rendered_html, base_url=_HTML_TEMPLATE_DIR).write_pdf(
        stylesheets=[CSS(filename=css_path)] if os.path.exists(css_path) else []
    )

    if not pdf_bytes or len(pdf_bytes) < 5000:
        raise RuntimeError(
            f"WeasyPrint produced a suspiciously small PDF ({len(pdf_bytes) if pdf_bytes else 0} bytes) "
            f"for template {template_filename}. Likely a rendering failure."
        )

    return pdf_bytes


# ── DOCX engine ───────────────────────────────────────────────────────────────

def merge_docx_file(docx_path: str, context: dict,
                    append_signature_tags: bool = False) -> bytes:
    """
    Fill the DOCX template at *docx_path* with *context* and return DOCX bytes.

    With append_signature_tags=True, appends Zoho Sign text tags so Zoho
    auto-creates a Signature field for each signer. Without them, Zoho's
    /submit returns error 9101 ("Add atleast one field for a signer").
    The rendered DOCX is then ready to upload to Zoho Sign (which converts
    to PDF server-side — no LibreOffice dependency).
    """
    tpl = DocxTemplate(docx_path)
    tpl.render(context)
    with tempfile.NamedTemporaryFile(suffix=".docx", delete=False) as tmp:
        tmp_path = tmp.name
    try:
        tpl.save(tmp_path)
        if append_signature_tags:
            _append_zoho_signature_tags(tmp_path)
        with open(tmp_path, "rb") as f:
            return f.read()
    finally:
        try:
            os.unlink(tmp_path)
        except OSError:
            pass


def docx_bytes_to_pdf(docx_bytes: bytes) -> bytes:
    """
    Convert rendered DOCX bytes to PDF via LibreOffice headless.

    Frappe Cloud's standard image does NOT include LibreOffice — signature
    letters avoid this path by uploading DOCX to Zoho directly. Plain-issue
    DOCX letters need LibreOffice on the bench; docx_to_pdf_bytes raises a
    clear error when it's missing.
    """
    with tempfile.NamedTemporaryFile(suffix=".docx", delete=False) as tmp:
        tmp_path = tmp.name
        tmp.write(docx_bytes)
    try:
        return docx_to_pdf_bytes(tmp_path)
    finally:
        try:
            os.unlink(tmp_path)
        except OSError:
            pass


def _append_zoho_signature_tags(docx_path: str) -> None:
    """
    Append a paragraph at the end of the DOCX with Zoho Sign text tags.

    Tags:
      {{S:R1*}}  — mandatory Signature field for recipient 1 (HR signatory)
      {{S:R2*}}  — mandatory Signature field for recipient 2 (letter recipient)

    Tags are rendered in white text so they are invisible in the rendered
    PDF (white-on-white) while Zoho still parses them from the document
    content (Zoho reads text, not color).

    Done as a post-processing step (not inside the template) to avoid the
    `{{ }}` syntax conflicting with docxtpl/Jinja2 expression evaluation.
    """
    from docx import Document
    from docx.shared import Pt, RGBColor

    doc = Document(docx_path)
    para = doc.add_paragraph()
    run = para.add_run("{{S:R1*}}                    {{S:R2*}}")
    run.font.size = Pt(8)
    run.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)  # white = invisible on white page
    doc.save(docx_path)


# ── Shared formatting helpers ─────────────────────────────────────────────────

def fmt_inr(value) -> str:
    """Format a number in Indian comma style, no currency symbol.
    e.g. 600000 → '6,00,000', 12345678 → '1,23,45,678'.
    """
    try:
        n = float(value or 0)
    except (TypeError, ValueError):
        return "0"
    s = "{:.0f}".format(n)
    if len(s) <= 3:
        return s
    last3, rest = s[-3:], s[:-3]
    grouped = ""
    while len(rest) > 2:
        grouped = "," + rest[-2:] + grouped
        rest = rest[:-2]
    return rest + grouped + "," + last3


def fmt_date(value) -> str:
    """Format a Frappe date field as '01 June 2026'."""
    if not value:
        return ""
    try:
        import frappe.utils
        return frappe.utils.formatdate(value, "dd MMMM yyyy")
    except Exception:
        return str(value)


def today_str() -> str:
    """Today's date as '01 June 2026'."""
    try:
        import frappe.utils
        return frappe.utils.formatdate(frappe.utils.today(), "dd MMMM yyyy")
    except Exception:
        return ""


def format_person_name(name) -> str:
    """
    Pretty-case a name string for letter display.

    Source data may arrive in mixed casing — sometimes all lowercase
    ("nalluri sudha"), sometimes all uppercase ("MOHD BALEEGH AHMED").
    For formal letters we want "Nalluri Sudha" / "Mohd Baleegh Ahmed".

    Strategy:
      - Empty / None → "" (caller decides how to handle missing)
      - Each word: if ALL-upper or ALL-lower → title-case it
                   if mixed-case (e.g., "McDonald", "O'Brien") → preserve as-is
      - Joined back with single spaces
    """
    if not name:
        return ""
    if not isinstance(name, str):
        name = str(name)
    words = name.strip().split()
    out = []
    for word in words:
        if not word:
            continue
        if word.isupper() or word.islower():
            out.append(word.capitalize())
        else:
            out.append(word)
    return " ".join(out)


def tenure_str(start_date, end_date) -> str:
    """Compute human-readable tenure: '2 years and 3 months'.
    end_date=None means up to today.
    """
    if not start_date:
        return ""
    try:
        from datetime import date, datetime
        if isinstance(start_date, str):
            start_date = datetime.strptime(start_date, "%Y-%m-%d").date()
        if end_date is None:
            end_date = date.today()
        elif isinstance(end_date, str):
            end_date = datetime.strptime(end_date, "%Y-%m-%d").date()

        years = end_date.year - start_date.year
        months = end_date.month - start_date.month
        if end_date.day < start_date.day:
            months -= 1
        if months < 0:
            years -= 1
            months += 12

        parts = []
        if years > 0:
            parts.append(f"{years} year{'s' if years != 1 else ''}")
        if months > 0:
            parts.append(f"{months} month{'s' if months != 1 else ''}")
        return " and ".join(parts) if parts else "less than a month"
    except Exception:
        return ""


@frappe.whitelist()
def health_check() -> dict:
    """
    HTTP-callable health check for the letter pipeline.

    Verifies: docxtpl import, WeasyPrint/libcairo availability, and the
    shipped HTML template directory.

    Call from any logged-in browser:
        https://<site>/api/method/globex_hr_letters.letters.merger.health_check
    """
    result = {
        "docxtpl_installed": False,
        "docxtpl_version": None,
        "html_template_dir": _HTML_TEMPLATE_DIR,
        "html_templates_found": [],
        "errors": [],
    }

    try:
        import docxtpl
        result["docxtpl_installed"] = True
        result["docxtpl_version"] = docxtpl.__version__
    except Exception as exc:
        result["errors"].append(f"docxtpl import: {exc!r}")

    try:
        result["html_templates_found"] = sorted(
            f for f in os.listdir(_HTML_TEMPLATE_DIR)
            if f.endswith(".html") and not f.startswith("_")
        )
    except Exception as exc:
        result["errors"].append(f"template dir: {exc!r}")

    # WeasyPrint / libcairo pre-flight
    try:
        from .pdf_check import check_pdf_dependencies
        result.update(check_pdf_dependencies())
    except Exception as exc:
        result["errors"].append(f"pdf_check: {exc!r}")

    return result
