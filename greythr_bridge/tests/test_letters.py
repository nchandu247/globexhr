"""
Tests for the DOCX letter merger.

These tests run without Frappe — they mock the frappe module.
"""
import os
import sys
import types
import unittest
from unittest.mock import MagicMock, patch


# ── Minimal frappe stub so merger.py can be imported offline ──────────────────
_frappe = types.ModuleType("frappe")
_frappe.throw = lambda msg, title=None: (_ for _ in ()).throw(ValueError(msg))
_frappe.utils = types.SimpleNamespace(
    formatdate=lambda v, fmt: str(v)
)
sys.modules.setdefault("frappe", _frappe)
sys.modules.setdefault("frappe.utils", _frappe.utils)

# Stub docxtpl
_docxtpl = types.ModuleType("docxtpl")
_docxtpl.DocxTemplate = MagicMock()
sys.modules.setdefault("docxtpl", _docxtpl)

from greythr_bridge.letters.merger import build_offer_context  # noqa: E402


class FakeDoc:
    """Minimal Job Offer document stub."""
    name = "OFR-2025-TEST"
    applicant_name = "Test Candidate"
    designation = "Software Engineer"
    department = "Engineering"
    offer_date = "2025-06-01"
    job_applicant = None
    applicant_email = "test.candidate@example.com"
    custom_title = "Mr."
    custom_date_of_joining = "2025-06-16"
    custom_annual_ctc = 600000
    custom_basic_monthly = 21600      # 50% of ~43200 gross
    custom_hra_monthly = 10800        # 50% of basic
    custom_conveyance_allowance_monthly = 1600
    custom_medical_allowance_monthly = 1250
    custom_special_allowance_monthly = 7950
    custom_employee_pf_monthly = 1800
    custom_employee_esi_monthly = 0
    custom_professional_tax_monthly = 200
    custom_employer_pf = 1950
    custom_employer_esiinsurance_monthly = 0
    custom_pf_opted_out = False
    custom_medical_insurance_opted_out = False
    custom_medical_insurance_annual_premium = 10000
    # Optional offer terms
    custom_work_location = "Hyderabad"
    custom_reporting_to = ""
    custom_probation_period = "6 months"
    custom_notice_period = "60 days"
    custom_joining_bonus = 50000
    custom_variable_pay_annual = 0
    custom_acceptance_deadline = "2025-06-10"


class TestBuildOfferContext(unittest.TestCase):

    def setUp(self):
        self.ctx = build_offer_context(FakeDoc())

    def test_candidate_name_present(self):
        self.assertEqual(self.ctx["candidate_name"], "Test Candidate")

    def test_title_flows_through(self):
        """custom_title field should be exposed as 'title' in context."""
        self.assertEqual(self.ctx["title"], "Mr.")

    def test_title_blank_when_field_missing(self):
        """If custom_title isn't set, title key should be empty string (not None)."""
        class NoTitleDoc:
            name = "OFR-NOTITLE"
            applicant_name = "X"
            designation = "Y"
            offer_date = "2025-06-01"
            custom_annual_ctc = 600000
            custom_basic_monthly = 21600
            custom_hra_monthly = 10800
            custom_conveyance_allowance_monthly = 1600
            custom_medical_allowance_monthly = 1250
            custom_special_allowance_monthly = 7950
            custom_employee_pf_monthly = 1800
            custom_employee_esi_monthly = 0
            custom_professional_tax_monthly = 200
            custom_employer_pf = 1950
            custom_employer_esiinsurance_monthly = 0
            # No custom_title

        ctx = build_offer_context(NoTitleDoc())
        self.assertEqual(ctx["title"], "")

    def test_gross_computed_from_components(self):
        expected_gross = 21600 + 10800 + 1600 + 1250 + 7950  # 43200
        self.assertEqual(self.ctx["gross_monthly_raw"], expected_gross)

    def test_net_computed_correctly(self):
        gross = 43200
        deductions = 1800 + 0 + 200  # pf + esi + pt
        self.assertEqual(self.ctx["net_take_home_raw"], gross - deductions)

    def test_inr_formatting_bare_no_symbol(self):
        # gross_monthly should be a bare number (no ₹) in Indian comma format
        self.assertNotIn("₹", self.ctx["gross_monthly"])
        self.assertEqual(self.ctx["gross_monthly"], "43,200")

    def test_indian_comma_formatting_for_lakhs(self):
        # 600000 → 6,00,000 (Indian style, not Western 600,000)
        self.assertEqual(self.ctx["annual_ctc"], "6,00,000")

    def test_esi_flag(self):
        # Gross 43200 > 21000 → ESI does not apply
        self.assertFalse(self.ctx["esi_applies"])

    def test_medical_insurance_included_when_esi_not_applicable(self):
        # ESI doesn't apply, medical not opted out → med_ins_monthly = 10000/12 = 833
        self.assertEqual(self.ctx["medical_insurance_monthly"], "833")

    def test_annual_values_computed_as_monthly_times_12(self):
        self.assertEqual(self.ctx["basic_annual"], "2,59,200")  # 21600 * 12
        self.assertEqual(self.ctx["hra_annual"], "1,29,600")    # 10800 * 12
        self.assertEqual(self.ctx["gross_annual"], "5,18,400")  # 43200 * 12

    def test_total_deductions(self):
        # emp_pf(1800) + emp_esi(0) + pt(200) = 2000 monthly, 24000 annually
        self.assertEqual(self.ctx["total_deductions_monthly"], "2,000")
        self.assertEqual(self.ctx["total_deductions_annual"], "24,000")

    def test_band_passthrough(self):
        # FakeDoc has no custom_band → default empty string
        self.assertEqual(self.ctx["band"], "")

    def test_all_required_keys_present(self):
        required = [
            # header / addressing
            "ref_number", "offer_date", "current_date", "title", "candidate_name",
            "candidate_email", "candidate_mobile", "candidate_address",
            "designation", "department", "band", "date_of_joining",
            "acceptance_deadline",
            # offer terms
            "work_location", "reporting_to", "probation_period", "notice_period",
            "joining_bonus", "variable_pay_annual",
            # CTC summary
            "annual_ctc", "monthly_ctc",
            "gross_monthly", "gross_annual",
            "net_take_home", "net_take_home_annual",
            # earnings monthly + annual
            "basic_monthly", "basic_annual",
            "hra_monthly", "hra_annual",
            "conveyance_monthly", "conveyance_annual",
            "medical_allowance_monthly", "medical_allowance_annual",
            "special_allowance_monthly", "special_allowance_annual",
            # employee deductions
            "employee_pf_monthly", "employee_pf_annual",
            "employee_esi_monthly", "employee_esi_annual",
            "professional_tax_monthly", "professional_tax_annual",
            "total_deductions_monthly", "total_deductions_annual",
            # employer contributions
            "employer_pf_monthly", "employer_pf_annual",
            "employer_esi_monthly", "employer_esi_annual",
            "medical_insurance_monthly", "medical_insurance_annual",
            "employer_deductions_annual",
            # flags
            "esi_applies", "pf_opted_out", "medical_opted_out",
            "has_joining_bonus", "has_variable_pay",
        ]
        for key in required:
            self.assertIn(key, self.ctx, f"Missing context key: {key}")

    def test_offer_terms_passthrough(self):
        self.assertEqual(self.ctx["work_location"], "Hyderabad")
        self.assertEqual(self.ctx["probation_period"], "6 months")
        self.assertEqual(self.ctx["notice_period"], "60 days")

    def test_joining_bonus_flag_set_when_amount_positive(self):
        self.assertTrue(self.ctx["has_joining_bonus"])
        self.assertEqual(self.ctx["joining_bonus"], "50,000")

    def test_variable_pay_flag_false_when_zero(self):
        self.assertFalse(self.ctx["has_variable_pay"])
        self.assertEqual(self.ctx["variable_pay_annual"], "0")

    def test_candidate_email_read_directly_from_job_offer(self):
        """Email should come from doc.applicant_email — no Job Applicant fetch needed."""
        self.assertEqual(self.ctx["candidate_email"], "test.candidate@example.com")

    def test_no_attribute_error_when_optional_fields_missing(self):
        """
        Regression test for 'JobOffer object has no attribute applicant'.
        Doc with no candidate-link / no opt-out flags should still produce
        a full context, not raise.
        """
        class MinimalDoc:
            name = "OFR-MIN"
            applicant_name = "Min Candidate"
            designation = "Eng"
            offer_date = "2025-06-01"
            custom_annual_ctc = 600000
            custom_basic_monthly = 21600
            custom_hra_monthly = 10800
            custom_conveyance_allowance_monthly = 1600
            custom_medical_allowance_monthly = 1250
            custom_special_allowance_monthly = 7950
            custom_employee_pf_monthly = 1800
            custom_employee_esi_monthly = 0
            custom_professional_tax_monthly = 200
            custom_employer_pf = 1950
            custom_employer_esiinsurance_monthly = 0
            # NO job_applicant, applicant_email, department, opt-out flags

        ctx = build_offer_context(MinimalDoc())
        self.assertEqual(ctx["candidate_email"], "")
        self.assertEqual(ctx["department"], "")
        self.assertFalse(ctx["pf_opted_out"])
        self.assertFalse(ctx["medical_opted_out"])

    def test_defaults_when_custom_fields_missing(self):
        """When optional custom fields aren't on the doc at all, defaults apply."""
        class BareDoc:
            name = "OFR-BARE"
            applicant_name = "Bare Candidate"
            designation = "Engineer"
            department = ""
            offer_date = "2025-06-01"
            job_applicant = None
            applicant_email = ""
            custom_annual_ctc = 600000
            custom_basic_monthly = 21600
            custom_hra_monthly = 10800
            custom_conveyance_allowance_monthly = 1600
            custom_medical_allowance_monthly = 1250
            custom_special_allowance_monthly = 7950
            custom_employee_pf_monthly = 1800
            custom_employee_esi_monthly = 0
            custom_professional_tax_monthly = 200
            custom_employer_pf = 1950
            custom_employer_esiinsurance_monthly = 0
            custom_pf_opted_out = False
            custom_medical_insurance_opted_out = False
            custom_medical_insurance_annual_premium = 10000
            # Note: NO optional offer-term attrs defined

        ctx = build_offer_context(BareDoc())
        self.assertEqual(ctx["work_location"], "Hyderabad")
        self.assertEqual(ctx["probation_period"], "6 months")
        self.assertEqual(ctx["notice_period"], "60 days")
        self.assertEqual(ctx["joining_bonus"], "0")
        self.assertFalse(ctx["has_joining_bonus"])
        self.assertEqual(ctx["reporting_to"], "")
        self.assertEqual(ctx["acceptance_deadline"], "")


class TestAppendZohoSignatureTags(unittest.TestCase):
    """
    Verify _append_zoho_signature_tags() injects Zoho text tags into a real
    DOCX so /submit doesn't fail with error 9101 'Add atleast one field for
    a signer'.
    """

    def test_appends_signature_tags_for_both_signers(self):
        import tempfile
        from docx import Document
        from greythr_bridge.letters.merger import _append_zoho_signature_tags

        with tempfile.NamedTemporaryFile(suffix=".docx", delete=False) as tmp:
            tmp_path = tmp.name
        try:
            doc = Document()
            doc.add_paragraph("Existing content")
            doc.save(tmp_path)

            _append_zoho_signature_tags(tmp_path)

            result = Document(tmp_path)
            all_text = "\n".join(p.text for p in result.paragraphs)
            self.assertIn("{{S:R1*}}", all_text,
                          "Signer 1 mandatory signature tag missing")
            self.assertIn("{{S:R2*}}", all_text,
                          "Signer 2 mandatory signature tag missing")
            self.assertIn("Existing content", all_text,
                          "Original content was destroyed")
        finally:
            os.unlink(tmp_path)

    def test_signature_tags_are_white_invisible(self):
        """
        Regression: tags were visible in black 8pt in the first signed PDF.
        Should now be white = invisible on white page (Zoho still parses text).
        """
        import tempfile
        from docx import Document
        from docx.shared import RGBColor
        from greythr_bridge.letters.merger import _append_zoho_signature_tags

        with tempfile.NamedTemporaryFile(suffix=".docx", delete=False) as tmp:
            tmp_path = tmp.name
        try:
            doc = Document()
            doc.save(tmp_path)
            _append_zoho_signature_tags(tmp_path)

            result = Document(tmp_path)
            # Find the paragraph containing the tags
            tag_para = None
            for p in result.paragraphs:
                if "{{S:R1*}}" in p.text:
                    tag_para = p
                    break
            self.assertIsNotNone(tag_para, "Tag paragraph not found")
            self.assertTrue(len(tag_para.runs) > 0, "Tag paragraph has no runs")
            color = tag_para.runs[0].font.color.rgb
            self.assertEqual(
                color, RGBColor(0xFF, 0xFF, 0xFF),
                f"Expected white (FFFFFF), got {color}",
            )
        finally:
            os.unlink(tmp_path)


class TestOfferContextDateOfJoining(unittest.TestCase):
    """Verify build_offer_context picks up custom_date_of_joining correctly."""

    def test_doj_truthy_when_set(self):
        """When DOJ is set on the doc, context's date_of_joining is non-empty."""
        ctx = build_offer_context(FakeDoc())
        # Note: actual format depends on frappe.utils.formatdate (not under test).
        # We just verify the field flows through to the context.
        self.assertTrue(
            ctx["date_of_joining"],
            "date_of_joining should be non-empty when custom_date_of_joining is set",
        )

    def test_doj_blank_when_missing(self):
        class NoDOJDoc:
            name = "OFR-NODOJ"
            applicant_name = "X"
            designation = "Y"
            offer_date = "2025-06-01"
            custom_annual_ctc = 600000
            custom_basic_monthly = 21600
            custom_hra_monthly = 10800
            custom_conveyance_allowance_monthly = 1600
            custom_medical_allowance_monthly = 1250
            custom_special_allowance_monthly = 7950
            custom_employee_pf_monthly = 1800
            custom_employee_esi_monthly = 0
            custom_professional_tax_monthly = 200
            custom_employer_pf = 1950
            custom_employer_esiinsurance_monthly = 0
            # No custom_date_of_joining attribute at all

        ctx = build_offer_context(NoDOJDoc())
        self.assertEqual(
            ctx["date_of_joining"], "",
            "date_of_joining should be empty string when field missing",
        )


class TestPdfCheck(unittest.TestCase):
    """Verify pdf_check.check_pdf_dependencies() returns expected structure."""

    def test_returns_all_expected_keys(self):
        from greythr_bridge.letters.pdf_check import check_pdf_dependencies
        result = check_pdf_dependencies()
        expected_keys = {
            "weasyprint_installed", "weasyprint_version",
            "libcairo_available", "libpango_available", "libgdk_pixbuf_available",
        }
        self.assertTrue(
            expected_keys.issubset(result.keys()),
            f"Missing keys: {expected_keys - set(result.keys())}",
        )
        # Booleans should be booleans
        for k in ("weasyprint_installed", "libcairo_available", "libpango_available", "libgdk_pixbuf_available"):
            self.assertIsInstance(result[k], bool, f"{k} should be bool")


def _weasyprint_can_render() -> bool:
    """
    Check whether WeasyPrint can actually render on this host.

    WeasyPrint imports fine on Windows but fails at render time because
    libgobject/libcairo are GTK libraries not present by default. On Frappe
    Cloud's Ubuntu these are pre-installed, so tests run there.

    This probes by attempting a minimal render. Returns True only if
    rendering actually works end-to-end.
    """
    try:
        from weasyprint import HTML
        HTML(string="<p>probe</p>").write_pdf()
        return True
    except Exception:
        return False


_WEASYPRINT_RENDER_OK = _weasyprint_can_render()


class TestMergeToPdfViaHtml(unittest.TestCase):
    """
    Tests for merge_to_pdf_via_html(). Skipped on hosts where WeasyPrint
    can't render (typically Windows dev machines missing GTK libs).
    These run on Frappe Cloud where libcairo/libpango/libgdk-pixbuf are present.
    """

    @unittest.skipUnless(_WEASYPRINT_RENDER_OK, "WeasyPrint cannot render on this host (missing GTK libs)")
    def test_returns_valid_pdf_bytes(self):
        from greythr_bridge.letters.merger import merge_to_pdf_via_html, build_offer_context
        ctx = build_offer_context(FakeDoc())

        pdf = merge_to_pdf_via_html("offer_letter.html", ctx)
        self.assertIsInstance(pdf, bytes)
        self.assertTrue(pdf.startswith(b"%PDF-"), "Output is not a valid PDF (missing magic bytes)")
        self.assertGreater(len(pdf), 10_000, "PDF unreasonably small — likely a rendering failure")

    @unittest.skipUnless(_WEASYPRINT_RENDER_OK, "WeasyPrint cannot render on this host")
    def test_pdf_contains_candidate_name(self):
        from greythr_bridge.letters.merger import merge_to_pdf_via_html, build_offer_context
        ctx = build_offer_context(FakeDoc())
        pdf = merge_to_pdf_via_html("offer_letter.html", ctx)
        # WeasyPrint embeds text such that ASCII strings appear in the PDF bytes
        self.assertIn(b"Test Candidate", pdf)

    @unittest.skipUnless(_WEASYPRINT_RENDER_OK, "WeasyPrint cannot render on this host")
    def test_pdf_includes_zoho_signature_tags(self):
        from greythr_bridge.letters.merger import merge_to_pdf_via_html, build_offer_context
        ctx = build_offer_context(FakeDoc())
        pdf = merge_to_pdf_via_html("offer_letter.html", ctx)
        # Tags rendered in white text but still present in PDF byte stream
        self.assertIn(b"S:R1*", pdf)
        self.assertIn(b"S:R2*", pdf)

    def test_raises_on_missing_template(self):
        """This test runs regardless of WeasyPrint render availability —
        the missing-template check happens before any rendering."""
        from greythr_bridge.letters.merger import merge_to_pdf_via_html
        with self.assertRaises(Exception):
            merge_to_pdf_via_html("nonexistent_template.html", {})


class TestBuildScriptPolish(unittest.TestCase):
    """
    Verify the polish passes in scripts/build_offer_template.py:
    - Fix 'Hi .{{ candidate_name }}' → 'Hi {{ candidate_name }}'
    - Insert reporting-manager paragraph after the joining sentence
    """

    def _make_fixture_docx(self, with_hi_dot=True, with_joining=True):
        """Create a tiny fixture DOCX containing the patterns we polish."""
        import tempfile
        from docx import Document
        with tempfile.NamedTemporaryFile(suffix=".docx", delete=False) as tmp:
            path = tmp.name
        doc = Document()
        if with_hi_dot:
            doc.add_paragraph("Hi .{{ candidate_name }},")
        doc.add_paragraph("Dear {{ candidate_name }},")
        if with_joining:
            doc.add_paragraph(
                "We are eager for you to join us as early as possible, "
                "ideally by {{ date_of_joining }}."
            )
        doc.add_paragraph("Yours sincerely,")
        doc.save(path)
        return path

    def test_fix_hi_dot_salutation(self):
        from docx import Document
        # Import the build module
        import importlib.util
        spec = importlib.util.spec_from_file_location(
            "build_offer_template",
            os.path.join(os.path.dirname(__file__), "..", "..", "scripts", "build_offer_template.py"),
        )
        mod = importlib.util.module_from_spec(spec)
        spec.loader.exec_module(mod)

        src = self._make_fixture_docx()
        dst = src.replace(".docx", "_out.docx")
        try:
            stats = mod.build(src, dst)
            self.assertGreaterEqual(stats["salutation_fixes"], 1,
                                    "Expected at least 1 'Hi .' fix")
            result = Document(dst)
            all_text = "\n".join(p.text for p in result.paragraphs)
            self.assertNotIn("Hi .{{", all_text, "'Hi .' pattern still present")
            self.assertIn("Hi {{ candidate_name }}", all_text,
                          "Fixed 'Hi ' salutation missing")
        finally:
            os.unlink(src)
            if os.path.exists(dst):
                os.unlink(dst)

    def test_inject_reporting_manager_line(self):
        from docx import Document
        import importlib.util
        spec = importlib.util.spec_from_file_location(
            "build_offer_template",
            os.path.join(os.path.dirname(__file__), "..", "..", "scripts", "build_offer_template.py"),
        )
        mod = importlib.util.module_from_spec(spec)
        spec.loader.exec_module(mod)

        src = self._make_fixture_docx()
        dst = src.replace(".docx", "_out.docx")
        try:
            stats = mod.build(src, dst)
            self.assertEqual(stats["reporting_manager_inserted"], 1,
                             "Expected exactly 1 reporting-manager line")
            result = Document(dst)
            all_text = "\n".join(p.text for p in result.paragraphs)
            self.assertIn("{% if reporting_to %}", all_text,
                          "Jinja conditional missing")
            self.assertIn("You will report to {{ reporting_to }}", all_text,
                          "Reporting manager sentence missing")
            self.assertIn("{% endif %}", all_text, "endif missing")
        finally:
            os.unlink(src)
            if os.path.exists(dst):
                os.unlink(dst)


# ─────────────────────────────────────────────────────────────────────────────
# Phase B tests — 6 new letter types + dispatcher + non_signing helper
# ─────────────────────────────────────────────────────────────────────────────


class _FakeJobOfferConsultant:
    name = "OFR-CONS-001"
    applicant_name = "Consultant Test"
    designation = "Senior Advisor"
    offer_date = "2026-06-01"
    job_applicant = None
    applicant_email = ""
    custom_offer_type = "Consultant"
    custom_engagement_duration_months = 6
    custom_professional_fees_monthly = 100000
    custom_date_of_joining = "2026-06-15"
    custom_title = "Mr."
    custom_work_location = "Hyderabad"
    custom_notice_period = "30 days"
    # Salary fields irrelevant for consultant — defaults
    custom_annual_ctc = 0
    custom_basic_monthly = 0
    custom_hra_monthly = 0
    custom_conveyance_allowance_monthly = 0
    custom_medical_allowance_monthly = 0
    custom_special_allowance_monthly = 0
    custom_employee_pf_monthly = 0
    custom_employee_esi_monthly = 0
    custom_professional_tax_monthly = 0
    custom_employer_pf = 0
    custom_employer_esiinsurance_monthly = 0


class _FakeJobOfferIntern:
    name = "OFR-INT-001"
    applicant_name = "Intern Test"
    designation = "Software Intern"
    offer_date = "2026-06-01"
    job_applicant = None
    applicant_email = ""
    custom_offer_type = "Intern"
    custom_internship_duration_months = 3
    custom_stipend_monthly = 15000
    custom_date_of_joining = "2026-06-15"
    custom_title = "Mr."
    custom_work_location = "Hyderabad"
    custom_reporting_to = ""
    # Other fields not used by intern context
    custom_annual_ctc = 0
    custom_basic_monthly = 0
    custom_hra_monthly = 0
    custom_conveyance_allowance_monthly = 0
    custom_medical_allowance_monthly = 0
    custom_special_allowance_monthly = 0
    custom_employee_pf_monthly = 0
    custom_employee_esi_monthly = 0
    custom_professional_tax_monthly = 0
    custom_employer_pf = 0
    custom_employer_esiinsurance_monthly = 0


class TestConsultantOfferContext(unittest.TestCase):
    def test_returns_consultant_specific_keys(self):
        from greythr_bridge.letters.merger import build_consultant_offer_context
        ctx = build_consultant_offer_context(_FakeJobOfferConsultant())
        for key in ("candidate_name", "designation", "engagement_start_date",
                    "engagement_duration", "professional_fees_monthly",
                    "professional_fees_annual", "work_location", "notice_period"):
            self.assertIn(key, ctx, f"Missing: {key}")
        self.assertEqual(ctx["candidate_name"], "Consultant Test")
        self.assertEqual(ctx["designation"], "Senior Advisor")
        # 6 months × 100000 = 1,00,000 monthly (formatted Indian comma)
        self.assertEqual(ctx["professional_fees_monthly"], "1,00,000")
        # Annual: 12 × 100000 = 12,00,000
        self.assertEqual(ctx["professional_fees_annual"], "12,00,000")

    def test_does_not_include_salary_table_keys(self):
        from greythr_bridge.letters.merger import build_consultant_offer_context
        ctx = build_consultant_offer_context(_FakeJobOfferConsultant())
        for forbidden in ("basic_monthly", "hra_monthly", "employee_pf_monthly",
                          "esi_applies", "total_deductions_annual"):
            self.assertNotIn(forbidden, ctx,
                             f"Consultant ctx should NOT include {forbidden}")


class TestInternOfferContext(unittest.TestCase):
    def test_returns_intern_specific_keys(self):
        from greythr_bridge.letters.merger import build_intern_offer_context
        ctx = build_intern_offer_context(_FakeJobOfferIntern())
        for key in ("candidate_name", "designation", "internship_start_date",
                    "internship_duration", "stipend_monthly", "stipend_total",
                    "work_location"):
            self.assertIn(key, ctx, f"Missing: {key}")
        self.assertEqual(ctx["stipend_monthly"], "15,000")
        # 3 months × 15000 = 45000
        self.assertEqual(ctx["stipend_total"], "45,000")

    def test_does_not_include_pf_or_salary_keys(self):
        from greythr_bridge.letters.merger import build_intern_offer_context
        ctx = build_intern_offer_context(_FakeJobOfferIntern())
        for forbidden in ("basic_monthly", "employee_pf_monthly",
                          "annual_ctc", "esi_applies"):
            self.assertNotIn(forbidden, ctx,
                             f"Intern ctx should NOT include {forbidden}")


class TestDispatcher(unittest.TestCase):
    """Verify dispatch_offer_letter() picks the correct template per offer_type."""

    def test_consultant_offer_type_routes_to_consultant_template(self):
        from greythr_bridge.letters.dispatch import dispatch_offer_letter
        tpl, ctx = dispatch_offer_letter(_FakeJobOfferConsultant())
        self.assertEqual(tpl, "consultant_offer_letter.html")
        self.assertEqual(ctx["candidate_name"], "Consultant Test")

    def test_intern_offer_type_routes_to_intern_template(self):
        from greythr_bridge.letters.dispatch import dispatch_offer_letter
        tpl, ctx = dispatch_offer_letter(_FakeJobOfferIntern())
        self.assertEqual(tpl, "intern_offer_letter.html")
        self.assertEqual(ctx["candidate_name"], "Intern Test")

    def test_full_time_routes_to_default_offer_template(self):
        from greythr_bridge.letters.dispatch import dispatch_offer_letter
        # FakeDoc is "Full-time" by default — no custom_offer_type set
        tpl, ctx = dispatch_offer_letter(FakeDoc())
        self.assertEqual(tpl, "offer_letter.html")

    def test_unknown_offer_type_falls_back_to_full_time(self):
        from greythr_bridge.letters.dispatch import dispatch_offer_letter

        class WeirdDoc:
            name = "WEIRD-001"
            applicant_name = "Weird Person"
            designation = "X"
            offer_date = "2026-06-01"
            job_applicant = None
            applicant_email = ""
            custom_offer_type = "Freelancer"  # not a known type
            custom_annual_ctc = 600000
            custom_basic_monthly = 21600
            custom_hra_monthly = 10800
            custom_conveyance_allowance_monthly = 1600
            custom_medical_allowance_monthly = 1250
            custom_special_allowance_monthly = 7950
            custom_employee_pf_monthly = 1800
            custom_employee_esi_monthly = 0
            custom_professional_tax_monthly = 200
            custom_employer_pf = 1950
            custom_employer_esiinsurance_monthly = 0

        tpl, ctx = dispatch_offer_letter(WeirdDoc())
        self.assertEqual(tpl, "offer_letter.html", "Unknown type should default to Full-time")


class TestPhaseBHTMLRendering(unittest.TestCase):
    """Verify each Phase B HTML template renders via Jinja2 without errors.
    This catches template syntax bugs without needing WeasyPrint (which can't
    render on Windows). Real PDF rendering happens on Frappe Cloud."""

    def _render(self, template, context):
        from jinja2 import Environment, FileSystemLoader, select_autoescape
        import greythr_bridge.letters.merger as m
        tmpl_dir = os.path.join(os.path.dirname(m.__file__), "..", "templates", "letters", "html")
        env = Environment(
            loader=FileSystemLoader(tmpl_dir),
            autoescape=select_autoescape(["html", "xml"]),
        )
        return env.get_template(template).render(**context)

    def test_consultant_offer_template_renders(self):
        from greythr_bridge.letters.merger import build_consultant_offer_context
        html = self._render("consultant_offer_letter.html",
                           build_consultant_offer_context(_FakeJobOfferConsultant()))
        self.assertIn("Consultant Test", html)
        self.assertIn("Senior Advisor", html)
        self.assertIn("Consultancy Engagement", html)
        self.assertIn("S:R1*", html)
        self.assertIn("S:R2*", html)

    def test_intern_offer_template_renders(self):
        from greythr_bridge.letters.merger import build_intern_offer_context
        html = self._render("intern_offer_letter.html",
                           build_intern_offer_context(_FakeJobOfferIntern()))
        self.assertIn("Intern Test", html)
        self.assertIn("Internship Offer", html)
        self.assertIn("S:R1*", html)
        self.assertIn("S:R2*", html)

    def test_increment_template_renders(self):
        # Use a synthesized context — testing builder requires frappe.get_doc
        context = {
            "ref_number": "SSA-001",
            "current_date": "01 June 2026",
            "employee_name": "Test Employee",
            "designation": "Engineer",
            "old_annual_ctc": "6,00,000",
            "new_annual_ctc": "7,50,000",
            "increment_amount": "1,50,000",
            "increment_percent": "25.0%",
            "effective_date": "01 June 2026",
            "previous_ssa": "SSA-PREV-001",
            "signatory_name": "Test Signatory",
        }
        html = self._render("increment_letter.html", context)
        self.assertIn("Test Employee", html)
        self.assertIn("Salary Increment", html)
        self.assertIn("1,50,000", html)
        self.assertIn("25.0%", html)

    def test_promotion_template_renders(self):
        context = {
            "ref_number": "EMP-001",
            "current_date": "01 June 2026",
            "employee_name": "Test Employee",
            "title": "",
            "old_designation": "Engineer",
            "new_designation": "Senior Engineer",
            "effective_date": "01 June 2026",
            "notes": "Outstanding performance",
            "signatory_name": "Test Signatory",
        }
        html = self._render("promotion_letter.html", context)
        self.assertIn("Senior Engineer", html)
        self.assertIn("Promotion", html)
        self.assertIn("Outstanding performance", html)

    def test_experience_template_renders(self):
        context = {
            "ref_number": "SEP-001",
            "current_date": "01 June 2026",
            "employee_name": "Test Employee",
            "designation": "Engineer",
            "date_of_joining": "01 January 2024",
            "last_working_day": "31 May 2026",
            "tenure": "2 years and 4 months",
            "signatory_name": "Test Signatory",
        }
        html = self._render("experience_letter.html", context)
        self.assertIn("Experience Letter", html)
        self.assertIn("To Whom It May Concern", html)
        self.assertIn("Test Employee", html)
        self.assertIn("2 years and 4 months", html)

    def test_relieving_template_renders(self):
        context = {
            "ref_number": "SEP-002",
            "current_date": "01 June 2026",
            "employee_name": "Test Employee",
            "designation": "Engineer",
            "date_of_joining": "01 January 2024",
            "last_working_day": "31 May 2026",
            "tenure": "2 years and 4 months",
            "signatory_name": "Test Signatory",
        }
        html = self._render("relieving_letter.html", context)
        self.assertIn("Relieving Letter", html)
        self.assertIn("Test Employee", html)

    def test_service_certificate_template_renders(self):
        context = {
            "ref_number": "EMP-002",
            "current_date": "01 June 2026",
            "employee_name": "Test Employee",
            "designation": "Engineer",
            "date_of_joining": "01 January 2024",
            "tenure_so_far": "2 years and 5 months",
            "signatory_name": "Test Signatory",
        }
        html = self._render("service_certificate.html", context)
        self.assertIn("Service Certificate", html)
        self.assertIn("Test Employee", html)
        self.assertIn("presently", html)
        self.assertIn("employed with Globex", html)

    # ── Bug 4 (2026-05-26): templates handle empty designation cleanly ───────

    def test_experience_template_handles_empty_designation_and_last_working_day(self):
        """For records where greytHR didn't return designation OR the
        Employee's relieving_date is still being set, the template must
        render without awkward gaps / dangling 'to' phrases."""
        context = {
            "ref_number": "GDS0021",
            "current_date": "26 May 2026",
            "employee_name": "Nalluri Sudha",
            "designation": "",          # empty
            "date_of_joining": "12 November 2017",
            "last_working_day": "",     # empty
            "tenure": "8 years and 6 months",
            "signatory_name": "Test",
        }
        html = self._render("experience_letter.html", context)
        # Must NOT contain the "served in the capacity of ." gap
        self.assertNotIn("served in the capacity of  .", html)
        self.assertNotIn("served in the capacity of  </strong>", html)
        # Must NOT contain dangling "to  ," (empty last_working_day)
        self.assertNotIn("from <strong>12 November 2017</strong> to <strong></strong>", html)
        # Should still mention tenure
        self.assertIn("8 years and 6 months", html)

    def test_relieving_template_handles_empty_designation_and_last_working_day(self):
        context = {
            "ref_number": "GDS0021",
            "current_date": "26 May 2026",
            "employee_name": "Nalluri Sudha",
            "designation": "",
            "date_of_joining": "12 November 2017",
            "last_working_day": "",
            "tenure": "8 years and 6 months",
            "signatory_name": "Test",
        }
        html = self._render("relieving_letter.html", context)
        # Must NOT contain "Nalluri Sudha, , has been relieved" (empty designation between commas)
        self.assertNotIn(", , has", html)
        # Must NOT contain "business hours on  ." (empty last_working_day)
        self.assertNotIn("business hours on <strong></strong>", html)
        self.assertIn("Nalluri Sudha", html)


class TestTenureCalculation(unittest.TestCase):
    """Test the _tenure() helper used by Experience/Relieving/Service Cert."""

    def test_two_years_three_months(self):
        from greythr_bridge.letters.merger import _tenure
        result = _tenure("2024-01-15", "2026-04-15")
        self.assertEqual(result, "2 years and 3 months")

    def test_one_year_one_month(self):
        from greythr_bridge.letters.merger import _tenure
        result = _tenure("2025-05-01", "2026-06-01")
        self.assertEqual(result, "1 year and 1 month")

    def test_exact_year_no_months(self):
        from greythr_bridge.letters.merger import _tenure
        result = _tenure("2025-06-01", "2026-06-01")
        self.assertEqual(result, "1 year")

    def test_less_than_month(self):
        from greythr_bridge.letters.merger import _tenure
        result = _tenure("2026-05-01", "2026-05-15")
        self.assertEqual(result, "less than a month")


class TestPhaseBCustomFields(unittest.TestCase):
    """Verify all 12 new Phase B custom fields are defined in fixtures."""

    def test_phase_b_fields_in_fixture(self):
        import json
        path = os.path.join(os.path.dirname(__file__), "..", "fixtures", "custom_field.json")
        with open(path) as f:
            data = json.load(f)
        fieldnames = {d["fieldname"] for d in data if "fieldname" in d}
        # 12 Phase B fields
        for expected in ("custom_offer_type", "custom_engagement_duration_months",
                          "custom_professional_fees_monthly", "custom_stipend_monthly",
                          "custom_internship_duration_months", "custom_annual_ctc",
                          "custom_send_increment_letter", "custom_increment_letter_generated",
                          "custom_send_experience_letter", "custom_send_relieving_letter",
                          "custom_promotion_letter_attached",
                          "custom_service_certificate_issued_at"):
            self.assertIn(expected, fieldnames, f"Missing Phase B field: {expected}")


class TestFormatName(unittest.TestCase):
    """The _format_name helper normalises greytHR's mixed-case name data
    for formal letter display. Added 2026-05-26 — greytHR returns names
    as "nalluri sudha" (lowercase) or "MOHD BALEEGH AHMED" (uppercase);
    we want letter output to read "Nalluri Sudha" / "Mohd Baleegh Ahmed"."""

    def test_lowercase_input_gets_title_cased(self):
        from greythr_bridge.letters.merger import _format_name
        self.assertEqual(_format_name("nalluri sudha"), "Nalluri Sudha")

    def test_uppercase_input_gets_title_cased(self):
        from greythr_bridge.letters.merger import _format_name
        self.assertEqual(_format_name("MOHD BALEEGH AHMED"), "Mohd Baleegh Ahmed")

    def test_already_title_cased_unchanged(self):
        from greythr_bridge.letters.merger import _format_name
        self.assertEqual(_format_name("Avinash Nalluri"), "Avinash Nalluri")

    def test_mixed_case_word_preserved(self):
        """Word containing both upper and lowercase letters (e.g., 'McDonald')
        must be preserved as-is, not title-cased into 'Mcdonald'."""
        from greythr_bridge.letters.merger import _format_name
        # McDonald has internal capital — must NOT be flattened
        self.assertEqual(_format_name("McDonald Smith"), "McDonald Smith")
        # Only McDonald is mixed; Smith is title-case already
        self.assertEqual(_format_name("McDonald SMITH"), "McDonald Smith")

    def test_empty_or_none_returns_empty_string(self):
        from greythr_bridge.letters.merger import _format_name
        self.assertEqual(_format_name(""), "")
        self.assertEqual(_format_name(None), "")
        self.assertEqual(_format_name("   "), "")

    def test_extra_whitespace_normalised_to_single_spaces(self):
        from greythr_bridge.letters.merger import _format_name
        self.assertEqual(_format_name("  nalluri    sudha  "), "Nalluri Sudha")


class TestSeparationContextBuilderFixes(unittest.TestCase):
    """Bugs 2 + 3 (2026-05-26): build_experience_context (also drives
    build_relieving_context) must use employee.name as ref_number and
    prefer Employee.relieving_date over Separation's date fields for
    last_working_day."""

    def _make_separation_with_employee(self, employee_attrs):
        """Build mock separation + mock employee, wire frappe.get_doc to
        return the employee."""
        sep = MagicMock()
        sep.name = "HR-EMP-SEP-2026-00001"
        sep.employee = "GDS0021"
        # Separation has NO relieving_date / boarding_end_date by default
        sep.relieving_date = None
        sep.boarding_end_date = None

        emp = MagicMock()
        emp.name = "GDS0021"
        for k, v in employee_attrs.items():
            setattr(emp, k, v)
        # Make MagicMock's getattr defaults match what merger.py reads
        emp.get = lambda field, *a, **kw: employee_attrs.get(field)

        return sep, emp

    def test_ref_number_uses_employee_name_not_separation_name(self):
        from greythr_bridge.letters import merger
        sep, emp = self._make_separation_with_employee({
            "employee_name": "nalluri sudha",
            "first_name": "nalluri sudha",
            "designation": "",
            "date_of_joining": "2017-11-12",
            "relieving_date": "2022-09-30",
        })
        with patch("frappe.get_doc", return_value=emp):
            context = merger.build_experience_context(sep)
        # Bug 2 fix: ref_number is the employee's GDS#### identifier
        self.assertEqual(context["ref_number"], "GDS0021")
        # NOT the separation's docname (which is what the bug used to show)
        self.assertNotEqual(context["ref_number"], "HR-EMP-SEP-2026-00001")

    def test_last_working_day_prefers_employee_relieving_date(self):
        """Bug 3 fix: Employee.relieving_date is Frappe HR's canonical
        last-working-day field. Must be preferred over Separation's
        boarding_end_date / relieving_date (which are often blank).
        We assert via fmt_date — when employee.relieving_date is set,
        fmt_date receives that value (proven via mock-call inspection)
        instead of being short-circuited to '' by the empty check."""
        from greythr_bridge.letters import merger
        sep, emp = self._make_separation_with_employee({
            "employee_name": "nalluri sudha",
            "first_name": "nalluri sudha",
            "designation": "",
            "date_of_joining": "2017-11-12",
            "relieving_date": "2022-09-30",  # SET on employee
        })
        # Patch fmt_date to a sentinel so we can verify it was called with
        # the employee's relieving_date (proving the new fallback chain
        # reached the right field). Before bug 3 fix, it would have been
        # called with None (Separation's empty fields) → returned "".
        captured = {}
        def _capture_fmt_date(v):
            if "relieving" not in captured:
                captured["relieving"] = v if v else None
            return str(v) if v else ""
        with patch("frappe.get_doc", return_value=emp), \
             patch.object(merger, "fmt_date", side_effect=lambda v: str(v) if v else ""):
            context = merger.build_experience_context(sep)
        # last_working_day must be truthy (the bug had it as "")
        self.assertTrue(context["last_working_day"],
            "last_working_day must NOT be empty — Bug 3's symptom was that "
            "this field was '' because the old fallback chain checked only "
            "separation.relieving_date / boarding_end_date (both blank).")
        # And specifically it should be the employee's relieving_date
        self.assertEqual(context["last_working_day"], "2022-09-30")

    def test_employee_name_run_through_format_name(self):
        """Bug 5 fix: greytHR lowercase 'nalluri sudha' renders as
        properly-cased 'Nalluri Sudha' in formal letters."""
        from greythr_bridge.letters import merger
        sep, emp = self._make_separation_with_employee({
            "employee_name": "nalluri sudha",
            "first_name": "nalluri sudha",
            "designation": "ENGINEER",
            "date_of_joining": "2017-11-12",
            "relieving_date": "2022-09-30",
        })
        with patch("frappe.get_doc", return_value=emp):
            context = merger.build_experience_context(sep)
        self.assertEqual(context["employee_name"], "Nalluri Sudha")
        # Same for designation
        self.assertEqual(context["designation"], "Engineer")


class TestSeparationLetterDualAttachment(unittest.TestCase):
    """Bug 1 (2026-05-26): separation letters attach to BOTH the Employee
    (primary, for permanent person-record) AND the Separation doc (secondary,
    for HR workflow view). Filename uses the Employee's GDS#### identifier
    regardless of which record it's attached to.

    Bug 1.5 (2026-05-26 v2): secondary attachment uses file_url LINKING
    instead of duplicating the bytes — Frappe was appending a hash suffix
    (e.g., GDS002170526e.pdf) when writing the same filename twice. Linking
    via file_url keeps a single physical file with the clean original name.
    """

    def _make_capturing_frappe_mock(self, primary_file_url="/private/files/Experience Letter - GDS0021.pdf"):
        """Build a frappe mock that captures File doc payloads + simulates
        Frappe assigning a file_url to the primary File on insert."""
        created_file_docs = []

        def _capture_get_doc(payload):
            mock = MagicMock()
            if isinstance(payload, dict) and payload.get("doctype") == "File":
                created_file_docs.append(payload)
                # Simulate Frappe assigning file_url to the primary File
                # (the one with content=bytes, NOT the one with file_url=...)
                if "content" in payload and "file_url" not in payload:
                    mock.file_url = primary_file_url
                else:
                    mock.file_url = payload.get("file_url")
            mock.insert = MagicMock()
            return mock

        return created_file_docs, _capture_get_doc

    def test_dual_attach_with_employee_filename(self):
        """generate_and_deliver creates 2 File rows when also_attach_to is
        passed; both use the file_name_suffix (the GDS#### identifier) in
        the filename instead of either attachment's docname."""
        from greythr_bridge.letters import non_signing
        created_file_docs, capture = self._make_capturing_frappe_mock()

        with patch("greythr_bridge.letters.non_signing.frappe") as frappe_mock:
            frappe_mock.get_doc.side_effect = capture
            with patch.object(non_signing, "merge_to_pdf_via_html",
                              return_value=b"x" * 10000):  # fake non-empty PDF
                non_signing.generate_and_deliver(
                    template_filename="experience_letter.html",
                    context={"ref_number": "GDS0021"},
                    attach_to=("Employee", "GDS0021"),
                    also_attach_to=("Employee Separation", "HR-EMP-SEP-2026-00001"),
                    file_label="Experience Letter",
                    file_name_suffix="GDS0021",
                    employee_doc=None,  # skip email
                )

        # Two File attachments created
        self.assertEqual(len(created_file_docs), 2)
        # Both use the same filename suffix (GDS0021), not the separation docname
        for fdoc in created_file_docs:
            self.assertIn("GDS0021", fdoc["file_name"])
            self.assertNotIn("HR-EMP-SEP", fdoc["file_name"])
        # Primary attaches to Employee
        primary = created_file_docs[0]
        self.assertEqual(primary["attached_to_doctype"], "Employee")
        self.assertEqual(primary["attached_to_name"], "GDS0021")
        # Secondary attaches to Employee Separation
        secondary = created_file_docs[1]
        self.assertEqual(secondary["attached_to_doctype"], "Employee Separation")
        self.assertEqual(secondary["attached_to_name"], "HR-EMP-SEP-2026-00001")

    def test_secondary_attachment_uses_file_url_linking_not_content_write(self):
        """Bug 1.5 v2: avoid Frappe's hash-suffix collision by linking
        the secondary File doc to the primary's file_url instead of
        writing duplicate content."""
        from greythr_bridge.letters import non_signing
        created_file_docs, capture = self._make_capturing_frappe_mock(
            primary_file_url="/private/files/Experience Letter - GDS0021.pdf"
        )

        with patch("greythr_bridge.letters.non_signing.frappe") as frappe_mock:
            frappe_mock.get_doc.side_effect = capture
            with patch.object(non_signing, "merge_to_pdf_via_html",
                              return_value=b"x" * 10000):
                non_signing.generate_and_deliver(
                    template_filename="experience_letter.html",
                    context={"ref_number": "GDS0021"},
                    attach_to=("Employee", "GDS0021"),
                    also_attach_to=("Employee Separation", "HR-EMP-SEP-2026-00001"),
                    file_label="Experience Letter",
                    file_name_suffix="GDS0021",
                    employee_doc=None,
                )

        primary, secondary = created_file_docs
        # Primary: WRITE mode — has content, no file_url
        self.assertIn("content", primary)
        self.assertNotIn("file_url", primary)
        # Secondary: LINK mode — has file_url pointing at primary's URL,
        # no duplicate content write
        self.assertNotIn("content", secondary)
        self.assertIn("file_url", secondary)
        self.assertEqual(secondary["file_url"],
                         "/private/files/Experience Letter - GDS0021.pdf")


class TestLastWorkingDayFallbackChain(unittest.TestCase):
    """Bug 3 v2 (2026-05-26): expanded fallback chain so we try MORE
    possible date fields on the Separation when Employee.relieving_date
    isn't set. Old chain only tried separation.relieving_date and
    boarding_end_date — many HR workflows fill in resignation_letter_date
    instead. New chain tries 4 fields."""

    def _make_emp_sep(self, employee_attrs, separation_attrs):
        sep = MagicMock()
        sep.name = "HR-EMP-SEP-2026-00001"
        sep.employee = "GDS0021"
        for k, v in separation_attrs.items():
            setattr(sep, k, v)
        # Also explicitly None-out fields not provided
        for f in ("relieving_date", "boarding_end_date", "resignation_letter_date"):
            if f not in separation_attrs:
                setattr(sep, f, None)

        emp = MagicMock()
        emp.name = "GDS0021"
        for k, v in employee_attrs.items():
            setattr(emp, k, v)
        emp.get = lambda f, *a, **kw: employee_attrs.get(f)
        return sep, emp

    def test_falls_through_to_boarding_end_date(self):
        from greythr_bridge.letters import merger
        sep, emp = self._make_emp_sep(
            employee_attrs={"employee_name": "X", "first_name": "X",
                            "designation": "", "date_of_joining": "2017-11-12",
                            "relieving_date": None},
            separation_attrs={"boarding_end_date": "2026-05-31"},
        )
        with patch("frappe.get_doc", return_value=emp), \
             patch.object(merger, "fmt_date", side_effect=lambda v: str(v) if v else ""):
            context = merger.build_experience_context(sep)
        self.assertEqual(context["last_working_day"], "2026-05-31")

    def test_falls_through_to_resignation_letter_date(self):
        from greythr_bridge.letters import merger
        sep, emp = self._make_emp_sep(
            employee_attrs={"employee_name": "X", "first_name": "X",
                            "designation": "", "date_of_joining": "2017-11-12",
                            "relieving_date": None},
            separation_attrs={"resignation_letter_date": "2026-04-15"},
        )
        with patch("frappe.get_doc", return_value=emp), \
             patch.object(merger, "fmt_date", side_effect=lambda v: str(v) if v else ""):
            context = merger.build_experience_context(sep)
        self.assertEqual(context["last_working_day"], "2026-04-15")

    def test_returns_empty_when_no_date_anywhere(self):
        """If none of the 4 fields are set, last_working_day is empty.
        The template's Jinja conditionals (Bug 4 fix) render cleanly
        without the date phrase."""
        from greythr_bridge.letters import merger
        sep, emp = self._make_emp_sep(
            employee_attrs={"employee_name": "X", "first_name": "X",
                            "designation": "", "date_of_joining": "2017-11-12",
                            "relieving_date": None},
            separation_attrs={},  # nothing set
        )
        with patch("frappe.get_doc", return_value=emp), \
             patch.object(merger, "fmt_date", side_effect=lambda v: str(v) if v else ""):
            context = merger.build_experience_context(sep)
        # Empty falsy value — template Jinja conditionals handle it
        self.assertFalse(context["last_working_day"])


if __name__ == "__main__":
    unittest.main()
